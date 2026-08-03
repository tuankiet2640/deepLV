"""Document translation orchestration tests.

Runs the real DocumentTranslator against a real SQLite database with a stub
provider, so the behaviour that only shows up end to end is covered: charging
each distinct string once, retrying, falling back to source text when a segment
cannot be translated, and storing a downloadable file.
"""

import io
import uuid

import pytest
import pytest_asyncio
from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from src.api.models import Base
from src.api.models.document_artifact import DocumentArtifact
from src.api.models.document_job import DocumentJob
from src.api.models.document_result import DocumentResult
from src.api.models.user import User
from src.api.services.document_translator import DocumentTranslator
from src.api.services.provider_manager import ResolvedProvider
from src.api.services.providers.base import ProviderError, TranslationProvider
from tests.document_fixtures import build_docx, build_pdf, build_pptx, build_xlsx


class StubProvider(TranslationProvider):
    """Records every call and can be told to fail for specific inputs."""

    provider_name = "stub"

    def __init__(self, fail_on: set[str] | None = None, fail_times: int = 99):
        self.calls: list[str] = []
        self.fail_on = fail_on or set()
        self.fail_times = fail_times
        self.failures: dict[str, int] = {}

    async def translate(self, text: str, source_lang: str, target_lang: str) -> str:
        self.calls.append(text)
        if text in self.fail_on:
            seen = self.failures.get(text, 0)
            if seen < self.fail_times:
                self.failures[text] = seen + 1
                raise ProviderError(self.provider_name, "boom")
        return f"[{target_lang}] {text}"

    async def validate_key(self) -> bool:
        return True


class StubProviderManager:
    """Stands in for ProviderManager without touching keys or credits."""

    def __init__(self, provider: TranslationProvider, used_own_key: bool = True):
        self.provider = provider
        self.used_own_key = used_own_key
        self.deducted: list[int] = []
        self.allow_credits = True

    async def resolve(self, provider_name, user, db, provider_key_id=None):
        return ResolvedProvider(provider=self.provider, used_own_key=self.used_own_key)

    async def deduct_credits(self, user, db, provider_name, char_count):
        self.deducted.append(char_count)
        return self.allow_credits


@pytest_asyncio.fixture
async def session():
    engine = create_async_engine("sqlite+aiosqlite:///:memory:")
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)

    factory = async_sessionmaker(engine, expire_on_commit=False)
    async with factory() as db:
        yield db
    await engine.dispose()


async def make_job(db, filename="report.docx", provider="stub", target="de"):
    user = User(
        id=uuid.uuid4(),
        email=f"{uuid.uuid4().hex}@example.com",
        password_hash="x",
        credits_balance=1000.0,
    )
    job = DocumentJob(
        id=uuid.uuid4(),
        user_id=user.id,
        original_filename=filename,
        file_size_bytes=1234,
        source_lang="en",
        target_lang=target,
        provider=provider,
        status="pending",
    )
    db.add_all([user, job])
    await db.commit()
    return job, user


async def fetch_artifact(db, job_id):
    result = await db.execute(select(DocumentArtifact).where(DocumentArtifact.job_id == job_id))
    return result.scalar_one_or_none()


async def fetch_result(db, job_id):
    result = await db.execute(select(DocumentResult).where(DocumentResult.job_id == job_id))
    return result.scalar_one_or_none()


async def test_translates_document_and_stores_a_downloadable_file(session):
    from docx import Document

    provider = StubProvider()
    translator = DocumentTranslator(StubProviderManager(provider))
    job, user = await make_job(session)

    await translator.translate_document(job, build_docx(), user, session)

    assert job.status == "completed"
    assert job.error_message is None
    assert job.completed_at is not None

    artifact = await fetch_artifact(session, job.id)
    assert artifact is not None
    assert artifact.filename == "report-de.docx"
    assert artifact.mime_type.endswith("wordprocessingml.document")
    assert artifact.format_preserved is True
    assert artifact.byte_size == len(artifact.content)
    assert artifact.failed_segments == 0
    assert artifact.translated_segments > 0

    # The stored bytes are a real, openable Word document with translated text.
    doc = Document(io.BytesIO(artifact.content))
    assert doc.paragraphs[0].text == "[de] Quarterly Report"
    assert doc.paragraphs[0].style.name == "Heading 1"


async def test_untranslatable_segments_are_never_sent_to_the_provider(session):
    provider = StubProvider()
    translator = DocumentTranslator(StubProviderManager(provider))
    job, user = await make_job(session)

    await translator.translate_document(job, build_docx(), user, session)

    assert "42" not in provider.calls
    assert "1,240,000" not in provider.calls
    assert not any("sales@example.com" == call for call in provider.calls)


async def test_repeated_text_is_translated_and_billed_once(session):
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    for row in range(1, 21):
        sheet.cell(row=row, column=1, value="Pending")
    sheet.cell(row=1, column=2, value="Unique label")
    buffer = io.BytesIO()
    workbook.save(buffer)

    provider = StubProvider()
    manager = StubProviderManager(provider, used_own_key=False)
    translator = DocumentTranslator(manager)
    job, user = await make_job(session, filename="status.xlsx", provider="openai")

    await translator.translate_document(job, buffer.getvalue(), user, session)

    # 20 identical cells cost one provider call and are billed once.
    assert provider.calls.count("Pending") == 1
    assert sorted(provider.calls) == ["Pending", "Unique label"]
    assert manager.deducted == [len("Pending") + len("Unique label")]

    # Every one of the 20 cells still gets the translation written into it.
    from openpyxl import load_workbook

    artifact = await fetch_artifact(session, job.id)
    sheet = load_workbook(io.BytesIO(artifact.content)).active
    assert [sheet.cell(row=r, column=1).value for r in range(1, 21)] == ["[de] Pending"] * 20


async def test_own_key_is_not_charged(session):
    manager = StubProviderManager(StubProvider(), used_own_key=True)
    translator = DocumentTranslator(manager)
    job, user = await make_job(session, filename="s.xlsx", provider="openai")

    await translator.translate_document(job, build_xlsx(), user, session)

    assert manager.deducted == []
    assert job.status == "completed"


async def test_insufficient_credits_fails_before_translating(session):
    provider = StubProvider()
    manager = StubProviderManager(provider, used_own_key=False)
    manager.allow_credits = False
    translator = DocumentTranslator(manager)
    job, user = await make_job(session, filename="s.xlsx", provider="openai")

    await translator.translate_document(job, build_xlsx(), user, session)

    assert job.status == "failed"
    assert "Insufficient credits" in job.error_message
    assert provider.calls == []
    assert await fetch_artifact(session, job.id) is None


async def test_transient_failure_is_retried(session):
    # Fails once, succeeds on the retry.
    provider = StubProvider(fail_on={"Widget"}, fail_times=1)
    translator = DocumentTranslator(StubProviderManager(provider))
    job, user = await make_job(session, filename="s.xlsx")

    await translator.translate_document(job, build_xlsx(), user, session)

    assert provider.calls.count("Widget") == 2
    assert job.status == "completed"
    assert job.error_message is None
    artifact = await fetch_artifact(session, job.id)
    assert artifact.failed_segments == 0


async def test_failed_segment_keeps_source_text_and_warns(session):
    from openpyxl import load_workbook

    provider = StubProvider(fail_on={"Widget"})
    translator = DocumentTranslator(StubProviderManager(provider))
    job, user = await make_job(session, filename="s.xlsx")

    await translator.translate_document(job, build_xlsx(), user, session)

    # The document is still delivered; the untranslated cell keeps its wording
    # rather than being filled with an error marker.
    assert job.status == "completed"
    assert "could not be translated" in job.error_message

    artifact = await fetch_artifact(session, job.id)
    assert artifact.failed_segments == 1
    sheet = load_workbook(io.BytesIO(artifact.content))["Sales"]
    assert sheet["A2"].value == "Widget"
    assert sheet["A1"].value == "[de] Product name"


async def test_total_failure_marks_the_job_failed(session):
    source = build_xlsx()
    all_texts = {"Product name", "Units sold", "Widget", "Total"}
    provider = StubProvider(fail_on=all_texts)
    translator = DocumentTranslator(StubProviderManager(provider))
    job, user = await make_job(session, filename="s.xlsx")

    await translator.translate_document(job, source, user, session)

    assert job.status == "failed"
    assert "all 4 text segments" in job.error_message
    assert await fetch_artifact(session, job.id) is None


async def test_unreadable_file_fails_the_job_with_a_clear_reason(session):
    translator = DocumentTranslator(StubProviderManager(StubProvider()))
    job, user = await make_job(session, filename="broken.docx")

    await translator.translate_document(job, b"this is not a docx", user, session)

    assert job.status == "failed"
    assert "Could not read DOCX" in job.error_message


async def test_plain_text_result_is_kept_for_previews(session):
    translator = DocumentTranslator(StubProviderManager(StubProvider()))
    job, user = await make_job(session)

    await translator.translate_document(job, build_docx(), user, session)

    result = await fetch_result(session, job.id)
    assert result is not None
    assert "[de] Quarterly Report" in result.translated_content
    assert result.chunk_count > 0
    assert result.total_characters > 0


async def test_long_segment_is_split_to_fit_provider_limits(session):
    from docx import Document

    sentence = "This is a sentence that repeats. "
    doc = Document()
    doc.add_paragraph(sentence * 200)  # ~6600 chars, over MAX_SEGMENT_CHARS
    buffer = io.BytesIO()
    doc.save(buffer)

    provider = StubProvider()
    translator = DocumentTranslator(StubProviderManager(provider))
    job, user = await make_job(session)

    await translator.translate_document(job, buffer.getvalue(), user, session)

    assert job.status == "completed"
    # One oversized segment became several provider calls, each within the limit.
    assert len(provider.calls) > 1
    assert all(len(call) <= 4000 for call in provider.calls)


async def test_pdf_job_produces_a_docx_artifact(session):
    translator = DocumentTranslator(StubProviderManager(StubProvider()))
    job, user = await make_job(session, filename="scan.pdf", target="fr")

    await translator.translate_document(
        job, build_pdf([["Page one text"], ["Page two text"]]), user, session
    )

    assert job.status == "completed"
    artifact = await fetch_artifact(session, job.id)
    assert artifact.filename == "scan-fr.docx"
    assert artifact.format_preserved is False


@pytest.mark.parametrize("filename", ["a.docx", "a.pptx", "a.xlsx", "a.txt", "a.html"])
async def test_every_in_place_format_reports_preserved_formatting(session, filename):
    builders = {
        "a.docx": build_docx,
        "a.pptx": build_pptx,
        "a.xlsx": build_xlsx,
        "a.txt": lambda: b"Hello world\n",
        "a.html": lambda: b"<html><body><p>Hello world</p></body></html>",
    }
    translator = DocumentTranslator(StubProviderManager(StubProvider()))
    job, user = await make_job(session, filename=filename)

    await translator.translate_document(job, builders[filename](), user, session)

    assert job.status == "completed", job.error_message
    artifact = await fetch_artifact(session, job.id)
    assert artifact.format_preserved is True
