"""Tests for format-preserving document translation output.

Covers the paragraph-level parsing/rebuild helpers in DocumentParser and the
DocumentTranslator orchestration that ties per-paragraph translation back
into a DOCX/PDF rebuild (or a graceful fallback to plain text).
"""

import io
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document
from pypdf import PdfReader

from src.api.models.document_job import DocumentJob
from src.api.models.user import User
from src.api.services.document_parser import DocumentParseError, DocumentParser
from src.api.services.document_translator import DocumentTranslator
from src.api.services.provider_manager import ResolvedProvider


def _make_docx_bytes(paragraphs: list[tuple[str, bool]]) -> bytes:
    """Build a DOCX with the given (text, bold) paragraphs."""
    doc = Document()
    for text, bold in paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(paragraph_text: str, table_rows: list[list[str]]) -> bytes:
    """Build a DOCX with one top-level paragraph followed by a table."""
    doc = Document()
    doc.add_paragraph(paragraph_text)
    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    for r, row in enumerate(table_rows):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_textbox(body_paragraph_text: str, textbox_paragraphs: list[str]) -> bytes:
    """Build a DOCX with a body paragraph plus a text box (``w:txbxContent``)
    containing its own paragraphs -- mirrors the callout/note-box shapes
    that live outside python-docx's own paragraph/table object model.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()
    doc.add_paragraph(body_paragraph_text)

    txbx_content = etree.SubElement(doc.element.body, qn("w:txbxContent"))
    for text in textbox_paragraphs:
        p = etree.SubElement(txbx_content, qn("w:p"))
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes(blocks: list[tuple[str, float, float, float, bool]]) -> bytes:
    """Build a single-page Letter PDF with text drawn at given positions.

    Each entry is (text, x, y, fontsize, bold); y is measured from the page
    bottom (reportlab's native coordinate system), so a larger y is higher
    up the page.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    for text, x, y, fontsize, bold in blocks:
        c.setFont("Helvetica-Bold" if bold else "Helvetica", fontsize)
        c.drawString(x, y, text)
    c.save()
    return buf.getvalue()


def _make_pdf_bytes_with_image() -> bytes:
    """Build a single-page PDF with a text block and one embedded image."""
    from PIL import Image
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    img = Image.new("RGB", (40, 40), color=(200, 30, 30))
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    img_buf.seek(0)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    c.setFont("Helvetica", 12)
    c.drawString(72, 700, "Text with a logo below.")
    c.drawImage(ImageReader(img_buf), 72, 600, width=40, height=40)
    c.save()
    return buf.getvalue()


class TestParseParagraphs:
    def test_docx_paragraphs_are_non_empty_and_in_order(self):
        original = _make_docx_bytes([("First", False), ("", False), ("Second", True)])
        parser = DocumentParser()

        paragraphs = parser.parse_paragraphs(original, "doc.docx")

        assert paragraphs == ["First", "Second"]

    def test_docx_table_cells_are_included_after_body_paragraphs(self):
        original = _make_docx_with_table(
            "Intro paragraph",
            [["Ten", "Nguyen Van A"], ["Chuc vu", "Ky su phan mem"]],
        )
        parser = DocumentParser()

        paragraphs = parser.parse_paragraphs(original, "doc.docx")

        assert paragraphs == [
            "Intro paragraph",
            "Ten",
            "Nguyen Van A",
            "Chuc vu",
            "Ky su phan mem",
        ]

    def test_docx_textbox_paragraphs_are_included(self):
        original = _make_docx_with_textbox(
            "Body paragraph",
            ["Luu y: note line one", "Note line two"],
        )
        parser = DocumentParser()

        paragraphs = parser.parse_paragraphs(original, "doc.docx")

        assert paragraphs == [
            "Body paragraph",
            "Luu y: note line one",
            "Note line two",
        ]

    def test_txt_splits_on_blank_lines(self):
        parser = DocumentParser()
        content = b"Para one.\n\nPara two."

        paragraphs = parser.parse_paragraphs(content, "file.txt")

        assert paragraphs == ["Para one.", "Para two."]

    def test_txt_without_blank_lines_is_single_paragraph(self):
        parser = DocumentParser()
        content = b"Just one block of text."

        paragraphs = parser.parse_paragraphs(content, "file.txt")

        assert paragraphs == ["Just one block of text."]

    def test_pdf_block_extraction_returns_bbox_font_and_bold(self):
        pdf_bytes = _make_pdf_bytes(
            [
                ("Bold Header", 72, 700, 14, True),
                ("A separate paragraph far below.", 72, 600, 10, False),
            ]
        )
        parser = DocumentParser()

        blocks = parser._pdf_text_blocks(pdf_bytes)

        assert len(blocks) == 2
        assert blocks[0].text == "Bold Header"
        assert blocks[0].bold is True
        assert blocks[0].font_size == pytest.approx(14, abs=0.5)
        assert blocks[0].page_index == 0
        x0, top, x1, bottom = blocks[0].bbox
        assert x0 < x1
        assert top < bottom
        assert blocks[1].text == "A separate paragraph far below."
        assert blocks[1].bold is False

    def test_pdf_close_lines_cluster_into_one_block(self):
        # Two lines of the same size/style with a small gap belong to the
        # same paragraph; a bold header right above them does not.
        pdf_bytes = _make_pdf_bytes(
            [
                ("Bold Header", 72, 700, 14, True),
                ("Body line one of a paragraph.", 72, 680, 10, False),
                ("Body line two, still same paragraph.", 72, 668, 10, False),
            ]
        )
        parser = DocumentParser()

        blocks = parser._pdf_text_blocks(pdf_bytes)

        assert len(blocks) == 2
        assert blocks[0].text == "Bold Header"
        assert blocks[1].text == (
            "Body line one of a paragraph.\nBody line two, still same paragraph."
        )

    def test_parse_paragraphs_pdf_returns_one_entry_per_block(self):
        pdf_bytes = _make_pdf_bytes(
            [
                ("First paragraph.", 72, 700, 10, False),
                ("Second paragraph, far below.", 72, 600, 10, False),
            ]
        )
        parser = DocumentParser()

        paragraphs = parser.parse_paragraphs(pdf_bytes, "doc.pdf")

        assert paragraphs == ["First paragraph.", "Second paragraph, far below."]

    def test_pdf_with_no_text_raises(self):
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=LETTER)
        c.save()
        empty_pdf = buf.getvalue()
        parser = DocumentParser()

        with pytest.raises(DocumentParseError):
            parser.parse_paragraphs(empty_pdf, "empty.pdf")


class TestBuildTranslatedDocx:
    def test_replaces_text_and_preserves_bold_formatting(self):
        original = _make_docx_bytes([("Xin chao", False), ("In dam", True)])
        parser = DocumentParser()

        rebuilt = parser.build_translated_docx(original, ["Hello", "Bold"])

        doc = Document(io.BytesIO(rebuilt))
        texts = [p.text for p in doc.paragraphs]
        bolds = [p.runs[0].bold if p.runs else None for p in doc.paragraphs]
        assert texts == ["Hello", "Bold"]
        assert bolds == [False, True]

    def test_extra_translated_paragraphs_are_ignored(self):
        original = _make_docx_bytes([("Only one", False)])
        parser = DocumentParser()

        # More translated entries than original non-empty paragraphs should
        # not raise or corrupt the document.
        rebuilt = parser.build_translated_docx(original, ["Translated", "Unused"])

        doc = Document(io.BytesIO(rebuilt))
        assert [p.text for p in doc.paragraphs] == ["Translated"]

    def test_table_cells_are_translated_not_left_in_original_language(self):
        original = _make_docx_with_table(
            "Intro paragraph",
            [["Ten", "Nguyen Van A"], ["Chuc vu", "Ky su phan mem"]],
        )
        parser = DocumentParser()
        translated = [
            "Introduction",
            "Name",
            "Nguyen Van A (translated)",
            "Position",
            "Software Engineer",
        ]

        rebuilt = parser.build_translated_docx(original, translated)

        doc = Document(io.BytesIO(rebuilt))
        assert doc.paragraphs[0].text == "Introduction"
        table = doc.tables[0]
        assert table.cell(0, 0).text == "Name"
        assert table.cell(0, 1).text == "Nguyen Van A (translated)"
        assert table.cell(1, 0).text == "Position"
        assert table.cell(1, 1).text == "Software Engineer"

    def test_textbox_paragraphs_are_translated_not_left_in_original_language(self):
        original = _make_docx_with_textbox(
            "Body paragraph",
            ["Luu y: note line one", "Note line two"],
        )
        parser = DocumentParser()
        translated = ["Body translated", "Note: translated line one", "Translated line two"]

        rebuilt = parser.build_translated_docx(original, translated)

        doc = Document(io.BytesIO(rebuilt))
        assert doc.paragraphs[0].text == "Body translated"
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        textbox_texts = [
            Paragraph(p_elem, doc).text
            for tb in doc.element.body.iter(qn("w:txbxContent"))
            for p_elem in tb.findall(qn("w:p"))
        ]
        assert textbox_texts == ["Note: translated line one", "Translated line two"]


class TestBuildTranslatedPdf:
    def test_rebuild_replaces_text(self):
        original = _make_pdf_bytes([("Hello world.", 72, 700, 12, False)])
        parser = DocumentParser()

        rebuilt = parser.build_translated_pdf(original, ["Xin chao the gioi."])

        assert rebuilt.startswith(b"%PDF-")
        reader = PdfReader(io.BytesIO(rebuilt))
        extracted = reader.pages[0].extract_text()
        assert "Xin chao the gioi." in extracted

    def test_rebuild_preserves_embedded_image(self):
        original = _make_pdf_bytes_with_image()
        parser = DocumentParser()
        orig_images = PdfReader(io.BytesIO(original)).pages[0].images
        assert len(orig_images) == 1

        rebuilt = parser.build_translated_pdf(original, ["Van ban voi logo ben duoi."])

        new_images = PdfReader(io.BytesIO(rebuilt)).pages[0].images
        assert len(new_images) == 1
        assert new_images[0].data == orig_images[0].data

    def test_vietnamese_text_renders_without_tofu_boxes(self):
        # Reportlab's default Helvetica only covers WinAnsi/Latin-1, so
        # Vietnamese diacritics used to come out as "�"/"■" tofu
        # boxes. build_translated_pdf must embed a Unicode-capable font.
        original = _make_pdf_bytes([("placeholder", 72, 700, 12, False)])
        parser = DocumentParser()
        vietnamese_text = "NGUYỄN THỊ THU UYÊN - Đinh Ngọc Mai, Vũ Thanh Hằng"

        rebuilt = parser.build_translated_pdf(original, [vietnamese_text])

        reader = PdfReader(io.BytesIO(rebuilt))
        extracted = reader.pages[0].extract_text()
        assert "�" not in extracted
        assert "■" not in extracted
        # Whitespace-normalized: the box may be narrow enough to word-wrap
        # the text across lines, which is fine -- only character loss/tofu
        # boxes are the regression this test guards against.
        assert " ".join(vietnamese_text.split()) in " ".join(extracted.split())

    def test_long_translated_text_does_not_crash(self):
        # A deliberately tiny original block -- translated text is much
        # longer than what the box could ever hold, even shrunk to the
        # font-size floor. Must degrade gracefully (shrink/overflow), not
        # raise or produce a corrupt PDF.
        original = _make_pdf_bytes([("Hi", 72, 700, 10, False)])
        parser = DocumentParser()
        long_text = "A very long translated sentence that is much longer than the original. " * 5

        rebuilt = parser.build_translated_pdf(original, [long_text])

        assert rebuilt.startswith(b"%PDF-")
        reader = PdfReader(io.BytesIO(rebuilt))
        extracted = reader.pages[0].extract_text()
        # The box is narrow enough that even individual words wrap across
        # lines at the font-size floor -- strip all whitespace so wrap
        # points (not just word boundaries) don't break the comparison;
        # this only checks no characters were lost or corrupted.
        assert "".join(long_text.split()) in "".join(extracted.split())

    def test_extra_translated_paragraphs_are_ignored(self):
        original = _make_pdf_bytes([("Only one.", 72, 700, 10, False)])
        parser = DocumentParser()

        # More translated entries than original blocks should not raise.
        rebuilt = parser.build_translated_pdf(original, ["Translated.", "Unused."])

        assert rebuilt.startswith(b"%PDF-")
        reader = PdfReader(io.BytesIO(rebuilt))
        extracted = reader.pages[0].extract_text()
        assert "Translated." in extracted


@pytest.fixture
def fake_provider():
    provider = MagicMock()
    provider.translate = AsyncMock(side_effect=lambda text, source_lang, target_lang: f"[{text}]")
    return provider


@pytest.fixture
def fake_provider_manager(fake_provider):
    manager = MagicMock()
    manager.resolve = AsyncMock(
        return_value=ResolvedProvider(provider=fake_provider, used_own_key=True)
    )
    manager.deduct_credits = AsyncMock(return_value=True)
    return manager


@pytest.fixture
def fake_db():
    db = MagicMock()
    db.commit = AsyncMock()
    db.added = []
    db.add = MagicMock(side_effect=lambda obj: db.added.append(obj))
    # translate_document() now looks up glossary terms for the job's
    # language pair -- no terms configured in these fixtures, so return an
    # empty result.
    execute_result = MagicMock()
    execute_result.scalars.return_value.all.return_value = []
    db.execute = AsyncMock(return_value=execute_result)
    return db


def _make_job(filename: str) -> DocumentJob:
    job = DocumentJob(
        original_filename=filename,
        file_size_bytes=100,
        source_lang="vi",
        target_lang="en",
        provider="openai",
        status="pending",
    )
    return job


class TestDocumentTranslator:
    async def test_docx_job_rebuilds_format_preserving_output(self, fake_provider_manager, fake_db):
        parser = DocumentParser()
        translator = DocumentTranslator(fake_provider_manager, parser)
        original = _make_docx_bytes([("Xin chao", False)])
        job = _make_job("greeting.docx")
        user = User(email="u@example.com", password_hash="x")

        await translator.translate_document(
            job=job,
            paragraphs=["Xin chao"],
            original_file_bytes=original,
            user=user,
            db=fake_db,
        )

        assert job.status == "completed"
        result = fake_db.added[0]
        assert result.output_format == "docx"
        assert result.translated_file_bytes is not None
        doc = Document(io.BytesIO(result.translated_file_bytes))
        assert doc.paragraphs[0].text == "[Xin chao]"

    async def test_txt_job_has_no_binary_output(self, fake_provider_manager, fake_db):
        parser = DocumentParser()
        translator = DocumentTranslator(fake_provider_manager, parser)
        job = _make_job("notes.txt")
        user = User(email="u@example.com", password_hash="x")

        await translator.translate_document(
            job=job,
            paragraphs=["Hello", "World"],
            original_file_bytes=b"Hello\n\nWorld",
            user=user,
            db=fake_db,
        )

        assert job.status == "completed"
        result = fake_db.added[0]
        assert result.output_format == "txt"
        assert result.translated_file_bytes is None
        assert result.translated_content == "[Hello]\n\n[World]"

    async def test_docx_rebuild_failure_falls_back_to_txt(self, fake_provider_manager, fake_db):
        parser = DocumentParser()
        # Not a real docx, so build_translated_docx will raise internally.
        translator = DocumentTranslator(fake_provider_manager, parser)
        job = _make_job("broken.docx")
        user = User(email="u@example.com", password_hash="x")

        await translator.translate_document(
            job=job,
            paragraphs=["Xin chao"],
            original_file_bytes=b"not a real docx",
            user=user,
            db=fake_db,
        )

        # Job should still complete; rebuild failure degrades gracefully.
        assert job.status == "completed"
        result = fake_db.added[0]
        assert result.output_format == "txt"
        assert result.translated_file_bytes is None

    async def test_one_paragraph_failure_does_not_fail_whole_job(
        self, fake_provider_manager, fake_db
    ):
        provider = fake_provider_manager.resolve.return_value.provider

        async def flaky_translate(text, source_lang, target_lang):
            if text == "bad":
                raise RuntimeError("provider exploded")
            return f"[{text}]"

        provider.translate = AsyncMock(side_effect=flaky_translate)

        parser = DocumentParser()
        translator = DocumentTranslator(fake_provider_manager, parser)
        job = _make_job("notes.txt")
        user = User(email="u@example.com", password_hash="x")

        await translator.translate_document(
            job=job,
            paragraphs=["good", "bad"],
            original_file_bytes=b"good\n\nbad",
            user=user,
            db=fake_db,
        )

        assert job.status == "completed"
        result = fake_db.added[0]
        assert "[good]" in result.translated_content
        assert "Translation error in section 2" in result.translated_content

    async def test_pdf_job_rebuilds_format_preserving_output(self, fake_provider_manager, fake_db):
        parser = DocumentParser()
        translator = DocumentTranslator(fake_provider_manager, parser)
        original = _make_pdf_bytes([("Xin chao", 72, 700, 12, False)])
        job = _make_job("greeting.pdf")
        user = User(email="u@example.com", password_hash="x")

        await translator.translate_document(
            job=job,
            paragraphs=["Xin chao"],
            original_file_bytes=original,
            user=user,
            db=fake_db,
        )

        assert job.status == "completed"
        result = fake_db.added[0]
        assert result.output_format == "pdf"
        assert result.translated_file_bytes is not None
        reader = PdfReader(io.BytesIO(result.translated_file_bytes))
        assert "[Xin chao]" in reader.pages[0].extract_text()

    async def test_pdf_rebuild_failure_falls_back_to_txt(self, fake_provider_manager, fake_db):
        parser = DocumentParser()
        # Not real PDF bytes, so build_translated_pdf will raise internally.
        translator = DocumentTranslator(fake_provider_manager, parser)
        job = _make_job("broken.pdf")
        user = User(email="u@example.com", password_hash="x")

        await translator.translate_document(
            job=job,
            paragraphs=["Xin chao"],
            original_file_bytes=b"not a real pdf",
            user=user,
            db=fake_db,
        )

        assert job.status == "completed"
        result = fake_db.added[0]
        assert result.output_format == "txt"
        assert result.translated_file_bytes is None
