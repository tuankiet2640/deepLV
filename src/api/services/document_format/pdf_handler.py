"""PDF handler — reads a PDF, writes a DOCX.

PDF has no editable text model: glyphs are placed at fixed coordinates with no
notion of a paragraph, and translated text is almost never the same length as the
source, so writing it back in place either overflows its box or leaves gaps.
Rewriting the page properly needs a layout engine that reflows and re-shapes the
whole page.

Rather than pretend, this handler extracts the reading order, translates it, and
emits a DOCX with page breaks where the PDF had them. Text, reading order and
page boundaries are preserved; the original visual layout is not, which is why
``preserves_format`` is False so the API can say so.
"""

import io
import re
from collections.abc import Iterator
from typing import Any

import structlog

from src.api.services.document_format.base import (
    DocumentFormatError,
    FormatHandler,
    TextUnit,
)

log = structlog.get_logger()

# Lines opening a list item start a new block even without a blank line before.
_BULLET_RE = re.compile(r"^\s*(?:[-\u2022\u00b7\u25cf\u25aa*]|\(?\d{1,3}[.)]|[a-z][.)])\s+")
_HYPHEN_BREAK_RE = re.compile(r"(\w)-$")


class _PdfDocument:
    """Extracted reading order: one list of text blocks per page."""

    def __init__(self, pages: list[list[str]]) -> None:
        self.pages = pages


class _BlockUnit:
    def __init__(self, document: _PdfDocument, page: int, block: int) -> None:
        self._document = document
        self._page = page
        self._block = block

    @property
    def context(self) -> str:
        return f"page {self._page + 1}"

    def get_text(self) -> str:
        return self._document.pages[self._page][self._block]

    def set_text(self, value: str) -> None:
        self._document.pages[self._page][self._block] = value


class PdfHandler(FormatHandler):
    extensions = ("pdf",)
    output_extension = ".docx"
    output_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    preserves_format = False

    def load(self, data: bytes) -> Any:
        from pypdf import PdfReader

        try:
            reader = PdfReader(io.BytesIO(data))
            pages = [_split_blocks(page.extract_text() or "") for page in reader.pages]
        except Exception as exc:
            raise DocumentFormatError(f"Could not read PDF file: {exc}") from exc

        if not any(pages):
            raise DocumentFormatError(
                "This PDF has no extractable text. Scanned or image-only PDFs need "
                "OCR before they can be translated."
            )
        log.info("pdf_extracted", pages=len(pages), blocks=sum(len(p) for p in pages))
        return _PdfDocument(pages)

    def units(self, document: Any) -> Iterator[TextUnit]:
        for page_index, blocks in enumerate(document.pages):
            for block_index in range(len(blocks)):
                yield _BlockUnit(document, page_index, block_index)

    def save(self, document: Any) -> bytes:
        from docx import Document

        out = Document()
        for page_index, blocks in enumerate(document.pages):
            for block_index, block in enumerate(blocks):
                paragraph = out.add_paragraph(block)
                # Marking the first block of each page rather than inserting an
                # explicit break keeps the output free of empty paragraphs.
                if page_index > 0 and block_index == 0:
                    paragraph.paragraph_format.page_break_before = True

        buffer = io.BytesIO()
        out.save(buffer)
        return buffer.getvalue()


def _split_blocks(text: str) -> list[str]:
    """Turn one page of extracted text into paragraph-sized blocks.

    Wrapped lines are rejoined so a sentence reaches the translator whole instead
    of arriving as fragments, which is the single biggest quality factor for PDF
    input. Blank lines and list markers end a block.
    """
    blocks: list[str] = []
    current: list[str] = []

    def flush() -> None:
        if current:
            joined = " ".join(current).strip()
            if joined:
                blocks.append(joined)
            current.clear()

    for raw_line in text.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            flush()
            continue
        if current and _BULLET_RE.match(line):
            flush()
        if current:
            # Join a word split across the line break without leaving a hyphen.
            previous = current[-1]
            if _HYPHEN_BREAK_RE.search(previous):
                current[-1] = previous[:-1] + line
                continue
        current.append(line)

    flush()
    return blocks
