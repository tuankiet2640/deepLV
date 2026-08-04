"""Document parsing service.

Extracts text content from PDF, DOCX, and TXT files,
then chunks it for translation.
"""

import io
from collections.abc import Iterator
from dataclasses import dataclass

import structlog

log = structlog.get_logger()

# Formats we can reconstruct as a translated file in their original format.
# Everything else (txt) falls back to a plain-text download.
REBUILDABLE_FORMATS = {"docx", "pdf"}

# Maximum characters per chunk for translation
MAX_CHUNK_SIZE = 5000

SUPPORTED_FORMATS = {"pdf", "docx", "txt"}

# A new PDF text block starts when the vertical gap to the previous line
# exceeds this multiple of the previous line's height.
PDF_BLOCK_GAP_MULTIPLIER = 1.5

# Shrink-to-fit bounds when redrawing translated text into a PDF block's
# original bounding box (translated text is rarely the same length as source).
PDF_MIN_FONT_SIZE = 6.0
PDF_FONT_SHRINK_STEP = 0.5


@dataclass(frozen=True)
class PdfTextBlock:
    """One paragraph-like text block extracted from a PDF page.

    Built by clustering pdfplumber's per-line extraction results by vertical
    proximity. parse_paragraphs() extracts these once (for translation and
    billing); build_translated_pdf() re-extracts the identical list, in the
    identical order, from the original bytes, then zips translated text back
    onto it by index -- the same lockstep pattern used for DOCX paragraphs.
    """

    page_index: int
    bbox: tuple[float, float, float, float]  # x0, top, x1, bottom (pdfplumber coords)
    text: str
    font_size: float
    bold: bool


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed."""

    pass


class DocumentParser:
    """Parses uploaded documents and extracts text content."""

    def parse(self, file_bytes: bytes, filename: str) -> str:
        """Parse a document and return its text content.

        Args:
            file_bytes: Raw bytes of the uploaded file.
            filename: Original filename (used to determine format).

        Returns:
            Extracted text content.

        Raises:
            DocumentParseError: If format is unsupported or parsing fails.
        """
        extension = self.get_extension(filename)

        if extension == "pdf":
            return self.parse_pdf(file_bytes)
        elif extension == "docx":
            return self.parse_docx(file_bytes)
        elif extension == "txt":
            return self.parse_txt(file_bytes)
        else:
            raise DocumentParseError(
                f"Unsupported file format: .{extension}. "
                f"Supported formats: {', '.join(sorted(SUPPORTED_FORMATS))}"
            )

    def parse_pdf(self, file_bytes: bytes) -> str:
        """Extract text from a PDF file.

        Args:
            file_bytes: Raw PDF bytes.

        Returns:
            Extracted text content.
        """
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

            content = "\n\n".join(pages)
            if not content.strip():
                raise DocumentParseError(
                    "PDF appears to contain no extractable text. "
                    "Scanned PDFs without OCR are not supported."
                )
            log.info("pdf_parsed", pages=len(reader.pages), chars=len(content))
            return content
        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse PDF: {e}") from e

    def parse_docx(self, file_bytes: bytes) -> str:
        """Extract text from a DOCX file.

        Args:
            file_bytes: Raw DOCX bytes.

        Returns:
            Extracted text content.
        """
        try:
            paragraphs = self._docx_paragraph_texts(file_bytes)
            content = "\n\n".join(paragraphs)
            if not content.strip():
                raise DocumentParseError("DOCX file contains no text content.")
            log.info("docx_parsed", paragraphs=len(paragraphs), chars=len(content))
            return content
        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse DOCX: {e}") from e

    def _docx_paragraph_texts(self, file_bytes: bytes) -> list[str]:
        """Return the non-empty paragraph texts of a DOCX, in document order.

        Includes body paragraphs followed by table cell paragraphs (tables
        are walked table-by-table, row-by-row, cell-by-cell). This exact
        order/filter is also used by ``build_translated_docx`` when walking
        the document to substitute translated text back in, so the two must
        stay in lockstep: the Nth entry returned here must always correspond
        to the Nth non-empty paragraph encountered when re-walking the
        document the same way.
        """
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        return list(self._iter_docx_paragraph_texts(doc))

    def _iter_docx_paragraph_texts(self, doc: object) -> Iterator[str]:
        """Yield non-empty paragraph texts: body paragraphs, table cells, then
        text box content (notes/callouts drawn as floating shapes)."""
        for para in doc.paragraphs:
            if para.text.strip():
                yield para.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            yield para.text
        for para in self._iter_docx_textbox_paragraphs(doc):
            if para.text.strip():
                yield para.text

    def _iter_docx_textbox_paragraphs(self, doc: object) -> Iterator[object]:
        """Yield python-docx Paragraph objects for every paragraph inside a
        text box (``w:txbxContent``), anywhere in the document body.

        Text boxes -- floating shapes with their own text, commonly used for
        callout/note boxes -- aren't exposed by python-docx's own object
        model (no ``doc.textboxes``), so this walks the underlying XML
        directly and wraps each raw paragraph element in a real ``Paragraph``
        so callers can read ``.text``/``.runs`` exactly like any other
        paragraph. Matches both the legacy VML (``v:textbox``) and modern
        DrawingML (``wps:txbx``) shape formats, since both wrap their text in
        a ``w:txbxContent`` element with ordinary ``w:p`` children -- only
        the tag being searched for matters, not what encloses it.
        """
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        for txbx_content in doc.element.body.iter(qn("w:txbxContent")):
            for p_element in txbx_content.findall(qn("w:p")):
                yield Paragraph(p_element, doc)

    def _pdf_text_blocks(self, file_bytes: bytes) -> list[PdfTextBlock]:
        """Extract paragraph-like text blocks from a PDF, page by page.

        Uses pdfplumber's per-line extraction (with bbox and font info per
        line), then clusters adjacent lines into blocks by vertical
        proximity -- pdfplumber has no built-in paragraph/block grouper.
        """
        import pdfplumber

        blocks: list[PdfTextBlock] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_index, page in enumerate(pdf.pages):
                lines = page.extract_text_lines()
                blocks.extend(self._cluster_pdf_lines(lines, page_index))
        return blocks

    def _line_style(self, line: dict) -> tuple[float, bool]:
        """Return a line's (font_size, bold) from its first character."""
        first_char = next(iter(line.get("chars", [])), None)
        if first_char is None:
            return (10.0, False)
        return (first_char["size"], "bold" in first_char["fontname"].lower())

    def _cluster_pdf_lines(self, lines: list[dict], page_index: int) -> list[PdfTextBlock]:
        """Group adjacent text lines into paragraph-like blocks.

        A new block starts whenever either:
        - the vertical gap to the previous line exceeds
          ``PDF_BLOCK_GAP_MULTIPLIER`` times the previous line's height, or
        - the font size or bold-state changes from the current block's
          style (a bold header immediately followed by body text, with no
          extra vertical gap, is still a different block -- a font change
          is a far more reliable boundary signal than gap distance alone).

        This is an approximation of paragraph boundaries, not exact layout
        analysis -- acceptable for the ~90% fidelity target. A block's own
        font_size/bold reflect its first line's style, same "keep the
        dominant run's formatting" precedent used for DOCX paragraphs.
        """
        blocks: list[PdfTextBlock] = []
        current_lines: list[dict] = []
        current_style: tuple[float, bool] = (10.0, False)

        def flush() -> None:
            if not current_lines:
                return
            text = "\n".join(line["text"] for line in current_lines if line["text"].strip())
            if text.strip():
                x0 = min(line["x0"] for line in current_lines)
                x1 = max(line["x1"] for line in current_lines)
                top = min(line["top"] for line in current_lines)
                bottom = max(line["bottom"] for line in current_lines)
                font_size, bold = current_style
                blocks.append(
                    PdfTextBlock(
                        page_index=page_index,
                        bbox=(x0, top, x1, bottom),
                        text=text,
                        font_size=font_size,
                        bold=bold,
                    )
                )
            current_lines.clear()

        previous: dict | None = None
        for line in lines:
            if not line["text"].strip():
                continue
            style = self._line_style(line)
            if current_lines and previous is not None:
                gap = line["top"] - previous["bottom"]
                line_height = max(previous["bottom"] - previous["top"], 1.0)
                if gap > PDF_BLOCK_GAP_MULTIPLIER * line_height or style != current_style:
                    flush()
            if not current_lines:
                current_style = style
            current_lines.append(line)
            previous = line
        flush()

        return blocks

    def parse_paragraphs(self, file_bytes: bytes, filename: str) -> list[str]:
        """Parse a document into its constituent paragraphs (or pages, for PDF).

        Unlike ``parse``, which returns one flattened string, this preserves
        the paragraph boundaries needed to rebuild a translated file in its
        original format (see ``build_translated_docx``/``build_translated_pdf``).

        Args:
            file_bytes: Raw bytes of the uploaded file.
            filename: Original filename (used to determine format).

        Returns:
            Non-empty paragraph/page texts, in document order.

        Raises:
            DocumentParseError: If format is unsupported or parsing fails.
        """
        extension = self.get_extension(filename)

        if extension == "docx":
            paragraphs = self._docx_paragraph_texts(file_bytes)
            if not paragraphs:
                raise DocumentParseError("DOCX file contains no text content.")
            return paragraphs
        elif extension == "pdf":
            blocks = self._pdf_text_blocks(file_bytes)
            paragraphs = [b.text for b in blocks]
            if not paragraphs:
                raise DocumentParseError(
                    "PDF appears to contain no extractable text. "
                    "Scanned PDFs without OCR are not supported."
                )
            return paragraphs
        elif extension == "txt":
            content = self.parse_txt(file_bytes)
            paragraphs = [p for p in content.split("\n\n") if p.strip()]
            return paragraphs or [content]
        else:
            raise DocumentParseError(
                f"Unsupported file format: .{extension}. "
                f"Supported formats: {', '.join(sorted(SUPPORTED_FORMATS))}"
            )

    def parse_txt(self, file_bytes: bytes) -> str:
        """Extract text from a plain text file.

        Args:
            file_bytes: Raw TXT bytes.

        Returns:
            Decoded text content.
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                content = file_bytes.decode("latin-1")

            if not content.strip():
                raise DocumentParseError("Text file is empty.")
            log.info("txt_parsed", chars=len(content))
            return content
        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Failed to parse text file: {e}") from e

    def build_translated_docx(
        self, original_file_bytes: bytes, translated_paragraphs: list[str]
    ) -> bytes:
        """Rebuild a DOCX with paragraph text replaced by its translation.

        Reopens the original document and walks its paragraphs, table
        cells, and text box content in the same order/filter as
        ``_docx_paragraph_texts``, substituting each non-empty paragraph's
        text with the corresponding translated entry. This keeps the
        original styles, headers/footers, and non-text paragraphs intact;
        only paragraph text changes.

        Args:
            original_file_bytes: The originally uploaded DOCX bytes.
            translated_paragraphs: Translated text, one entry per non-empty
                paragraph, in the same order ``_docx_paragraph_texts`` returns.

        Returns:
            Bytes of the rebuilt DOCX file.
        """
        from docx import Document

        doc = Document(io.BytesIO(original_file_bytes))
        index = 0

        def apply(para: object) -> None:
            nonlocal index
            if not para.text.strip():
                return
            if index < len(translated_paragraphs):
                self._set_paragraph_text(para, translated_paragraphs[index])
            index += 1

        for para in doc.paragraphs:
            apply(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        apply(para)
        for textbox_para in self._iter_docx_textbox_paragraphs(doc):
            apply(textbox_para)

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    def _set_paragraph_text(self, paragraph: object, new_text: str) -> None:
        """Replace a paragraph's visible text while preserving its formatting.

        Keeps the first run (and its font/bold/italic/etc.) and empties any
        remaining runs, rather than rebuilding runs from scratch, so
        character-level formatting on the paragraph's dominant run survives.
        """
        runs = paragraph.runs
        if not runs:
            paragraph.add_run(new_text)
            return
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""

    def build_translated_pdf(
        self, original_file_bytes: bytes, translated_paragraphs: list[str]
    ) -> bytes:
        """Rebuild a PDF in place: occlude each original text block, then
        redraw translated text into the same bounding box.

        Preserves images, lines, backgrounds, and page layout, since only
        the detected text-block regions are painted over and redrawn --
        everything else on the page is untouched. Not pixel-perfect (see
        limitations below), but targets ~90% visual fidelity versus the
        prior from-scratch approach, which discarded all layout and images.

        Technique: for each page, build a same-size reportlab overlay page
        that (1) paints an opaque white rectangle over every original text
        block's bbox, then (2) draws the translated text into that same
        bbox, shrinking font size to fit if needed. That overlay is then
        composited onto the original page via ``pypdf``'s ``merge_page``.
        Note this is occlusion, not true redaction: the original glyphs are
        still technically present in the PDF underneath the white
        rectangle, just visually covered -- fine for a translation tool, not
        a substitute for a security/redaction tool.

        Known fidelity limitations:
        - Font family always becomes DejaVu Sans (regular/bold); italic and
          the original font family are not preserved, only a bold/non-bold
          distinction.
        - The occlusion fill is always opaque white, so text sitting on a
          non-white background (shaded cells, colored boxes) will show a
          white patch behind the translated text.
        - Original text alignment (center/right/justify) isn't detected;
          redraw is always left-aligned within the original bbox.
        - Block clustering approximates paragraph boundaries by vertical
          gap and font/bold changes; it is not exact layout analysis, and
          rotated pages or complex multi-column layouts are out of scope.
        - If translated text is far longer than the original block even at
          the minimum font size, the overflow is drawn past the box's
          original bottom edge rather than clipped or raising an error.

        Args:
            original_file_bytes: The originally uploaded PDF bytes.
            translated_paragraphs: Translated text, one entry per text
                block, in the same order ``_pdf_text_blocks`` returns.

        Returns:
            Bytes of the rebuilt PDF file.
        """
        from pypdf import PdfReader, PdfWriter

        self._register_unicode_font()

        blocks = self._pdf_text_blocks(original_file_bytes)
        pages_to_blocks: dict[int, list[tuple[PdfTextBlock, str]]] = {}
        for block, translated_text in zip(blocks, translated_paragraphs, strict=False):
            pages_to_blocks.setdefault(block.page_index, []).append((block, translated_text))

        # Clone into the writer up front (rather than merging onto pages
        # still owned by a bare PdfReader) so every page is already
        # attached to the writer before merge_page runs on it.
        writer = PdfWriter(clone_from=io.BytesIO(original_file_bytes))

        for page_index, page in enumerate(writer.pages):
            page_blocks = pages_to_blocks.get(page_index)
            if page_blocks:
                page_width = float(page.mediabox.width)
                page_height = float(page.mediabox.height)
                overlay_bytes = self._build_pdf_overlay_page(page_width, page_height, page_blocks)
                overlay_reader = PdfReader(io.BytesIO(overlay_bytes))
                page.merge_page(overlay_reader.pages[0])

        out = io.BytesIO()
        writer.write(out)
        return out.getvalue()

    def _build_pdf_overlay_page(
        self,
        page_width: float,
        page_height: float,
        page_blocks: list[tuple["PdfTextBlock", str]],
    ) -> bytes:
        """Build a single-page PDF: white rectangles over every block's bbox,
        then the translated text drawn into those same boxes.

        Both passes happen on one reportlab canvas page, in draw order, so
        every occlusion rectangle (pass 1) is guaranteed to be painted
        before any translated text (pass 2) -- otherwise a later block's
        occlusion could paint over an earlier block's freshly-drawn text.
        """
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=(page_width, page_height))

        c.setFillColorRGB(1, 1, 1)
        for block, _translated_text in page_blocks:
            x0, top, x1, bottom = block.bbox
            c.rect(x0, page_height - bottom, x1 - x0, bottom - top, fill=1, stroke=0)

        for block, translated_text in page_blocks:
            self._draw_pdf_block_text(c, block, translated_text, page_height)

        c.save()
        return buf.getvalue()

    def _draw_pdf_block_text(
        self, c: object, block: "PdfTextBlock", text: str, page_height: float
    ) -> None:
        """Draw translated text into a block's original bbox, shrinking font
        size to fit -- translated text is rarely the same length as source."""
        from xml.sax.saxutils import escape

        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import Paragraph

        x0, top, x1, bottom = block.bbox
        height = bottom - top
        font_name = "DejaVuSans-Bold" if block.bold else "DejaVuSans"
        font_size = block.font_size or 10.0
        # Reportlab's Paragraph interprets its text as a small XML dialect;
        # escape it so literal &, <, > in translated text don't break rendering.
        safe_text = escape(text).replace("\n", "<br/>")

        # Wrapping to the raw block width degenerates into a single
        # character per line for very narrow blocks (e.g. a page number),
        # producing an ugly vertical letter-stack that can run down the
        # whole page. Floor the wrap width at ~4 characters' worth so
        # narrow blocks instead overflow a bit to the right -- a much
        # smaller visual cost than an unbounded vertical explosion.
        wrap_width = max(x1 - x0, font_size * 4)

        para = Paragraph(safe_text, ParagraphStyle("b", fontName=font_name, fontSize=font_size))
        _, required_height = para.wrap(wrap_width, height)
        while required_height > height and font_size > PDF_MIN_FONT_SIZE:
            font_size = max(PDF_MIN_FONT_SIZE, font_size - PDF_FONT_SHRINK_STEP)
            wrap_width = max(x1 - x0, font_size * 4)
            style = ParagraphStyle("b", fontName=font_name, fontSize=font_size)
            para = Paragraph(safe_text, style)
            _, required_height = para.wrap(wrap_width, height)

        y = page_height - top - required_height
        para.drawOn(c, x0, y)

    _unicode_font_registered = False

    def _register_unicode_font(self) -> None:
        """Register the vendored DejaVu Sans font with reportlab, once."""
        if DocumentParser._unicode_font_registered:
            return

        from pathlib import Path

        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        fonts_dir = Path(__file__).resolve().parent.parent / "assets" / "fonts"
        pdfmetrics.registerFont(TTFont("DejaVuSans", str(fonts_dir / "DejaVuSans.ttf")))
        pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", str(fonts_dir / "DejaVuSans-Bold.ttf")))
        pdfmetrics.registerFontFamily("DejaVuSans", normal="DejaVuSans", bold="DejaVuSans-Bold")
        DocumentParser._unicode_font_registered = True

    def chunk_text(self, text: str, max_chunk_size: int = MAX_CHUNK_SIZE) -> list[str]:
        """Split text into chunks suitable for translation.

        Splits on paragraph boundaries where possible, falling back
        to sentence boundaries, then hard splits at max_chunk_size.

        Args:
            text: Full text to chunk.
            max_chunk_size: Maximum characters per chunk.

        Returns:
            List of text chunks.
        """
        if len(text) <= max_chunk_size:
            return [text]

        chunks: list[str] = []
        # Split by paragraphs first
        paragraphs = text.split("\n\n")
        current_chunk = ""

        for paragraph in paragraphs:
            # If a single paragraph exceeds max size, split it further
            if len(paragraph) > max_chunk_size:
                # Flush current chunk
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                    current_chunk = ""

                # Split long paragraph by sentences
                sentences = self._split_sentences(paragraph)
                for sentence in sentences:
                    # If a single sentence exceeds max size, hard-split it
                    if len(sentence) > max_chunk_size:
                        if current_chunk.strip():
                            chunks.append(current_chunk.strip())
                            current_chunk = ""
                        chunks.extend(self._hard_split(sentence, max_chunk_size))
                    elif len(current_chunk) + len(sentence) + 1 > max_chunk_size:
                        if current_chunk.strip():
                            chunks.append(current_chunk.strip())
                        current_chunk = sentence
                    else:
                        current_chunk = (
                            current_chunk + " " + sentence if current_chunk else sentence
                        )

            elif len(current_chunk) + len(paragraph) + 2 > max_chunk_size:
                # Current chunk would exceed limit, start a new one
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph

        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        log.info("text_chunked", total_chars=len(text), chunks=len(chunks))
        return chunks

    def _hard_split(self, text: str, max_chunk_size: int) -> list[str]:
        """Hard-split text into chunks of max_chunk_size when no natural boundaries exist."""
        chunks = []
        for i in range(0, len(text), max_chunk_size):
            chunk = text[i : i + max_chunk_size]
            if chunk.strip():
                chunks.append(chunk.strip())
        return chunks

    def _split_sentences(self, text: str) -> list[str]:
        """Split text into sentences (simple heuristic)."""
        import re

        sentences = re.split(r"(?<=[.!?])\s+", text)
        return [s for s in sentences if s.strip()]

    def get_extension(self, filename: str) -> str:
        """Get lowercase file extension without dot."""
        if "." not in filename:
            return ""
        return filename.rsplit(".", 1)[-1].lower()
